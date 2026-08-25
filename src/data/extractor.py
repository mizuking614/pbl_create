import re
import sys
from pathlib import Path
from PIL import Image
import pypdf
import pypdfium2
import pytesseract
import docx

def read_text_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            return path.read_text(encoding="cp932")
        except Exception:
            return ""

def read_pdf(path: Path) -> str:
    try:
        reader = pypdf.PdfReader(str(path))
        text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
        
        # Fallback to OCR if less than 50 characters extracted
        if len(text.strip()) < 50:
            text = read_pdf_with_ocr(path)
        return text
    except Exception as e:
        print(f"PDF read error: {path.name}: {e}", file=sys.stderr)
        return ""

def read_pdf_page_with_ocr(page_img: Image.Image, lang: str = "jpn+eng") -> str:
    try:
        return pytesseract.image_to_string(page_img, lang=lang)
    except Exception as e:
        print(f"pytesseract error: {e}", file=sys.stderr)
        return ""

def read_pdf_with_ocr(path: Path, lang: str = "jpn+eng") -> str:
    text = ""
    try:
        pdf = pypdfium2.PdfDocument(str(path))
        for page in pdf:
            bitmap = page.render(scale=2) # Render at 150 DPI equivalent
            pil_img = bitmap.to_pil()
            page_text = read_pdf_page_with_ocr(pil_img, lang=lang)
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        print(f"PDF OCR render error: {path.name}: {e}", file=sys.stderr)
    return text

def read_image(path: Path, lang: str = "jpn+eng") -> str:
    try:
        with Image.open(path) as img:
            return pytesseract.image_to_string(img, lang=lang)
    except Exception as e:
        print(f"Image OCR error: {path.name}: {e}", file=sys.stderr)
        return ""

def read_docx(path: Path) -> str:
    try:
        doc = docx.Document(str(path))
        text = []
        for p in doc.paragraphs:
            if p.text:
                text.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text]
                if row_text:
                    text.append(" | ".join(row_text))
        return "\n".join(text)
    except Exception as e:
        print(f"Docx read error: {path.name}: {e}", file=sys.stderr)
        return ""

def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".py", ".csv", ".json"}:
        return read_text_file(path)
    elif suffix == ".pdf":
        return read_pdf(path)
    elif suffix in {".docx", ".doc"}:
        return read_docx(path)
    elif suffix in {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif"}:
        return read_image(path)
    return ""

def tokenize(text: str) -> list[str]:
    # Basic word tokenization (removing punctuation and splitting on space)
    # Handles both English words and basic Japanese boundaries
    cleaned = re.sub(r'[^\w\s\u3040-\u309f\u30a0-\u30ff\u4e00-\u9faf]', ' ', text)
    words = cleaned.split()
    return [w.lower() for w in words if len(w) > 1]

def split_sentences(text: str) -> list[str]:
    # Splits by Japanese period (。) or English period followed by space
    sentences = re.split(r'(?<=。)|(?<=\.)\s+', text)
    return [s.strip() for s in sentences if s.strip()]
