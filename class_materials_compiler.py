from __future__ import annotations

import argparse
import html
import os
import json
import math
import re
import shutil
import sys
import subprocess
import tempfile
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable
from urllib.parse import quote

# Automatic installation of reportlab for summary PDF generation
try:
    import reportlab
except ImportError:
    print("Installing reportlab for PDF generation...", file=sys.stderr)
    try:
        subprocess.run([sys.executable, "-m", "pip", "install", "reportlab"], check=True, stdout=subprocess.DEVNULL)
    except Exception as e:
        print(f"Warning: Failed to install reportlab: {e}", file=sys.stderr)


APP_DIR = ".class_materials"
CONFIG_FILE = "courses.json"
SUMMARIES_FILE = "summaries.json"
QUESTIONS_FILE = "questions.json"
DEFAULT_MODEL = "gpt-5.2"
MAX_AI_CHARS = 18000
SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".csv",
    ".pdf",
    ".png",
    ".jpg",
    ".jpeg",
    ".bmp",
    ".tif",
    ".tiff",
    ".docx",
    ".doc",
}
JAPANESE_STOPWORDS = {
    "これ",
    "それ",
    "ため",
    "もの",
    "こと",
    "ここ",
    "そこ",
    "する",
    "ある",
    "いる",
    "です",
    "ます",
    "から",
    "まで",
    "について",
}
ENGLISH_STOPWORDS = {
    "and",
    "are",
    "for",
    "from",
    "into",
    "that",
    "the",
    "this",
    "with",
    "you",
    "your",
}


@dataclass
class LinkRecord:
    title: str
    url: str
    memo: str = ""


@dataclass
class AttendanceRecord:
    date: str
    class_round: int
    status: str
    memo: str


@dataclass
class Course:
    name: str
    teacher: str
    keywords: list[str]
    folder: str
    learned_terms: list[str]
    attendance: list[AttendanceRecord]
    links: list[LinkRecord] = field(default_factory=list)


@dataclass
class MaterialRecord:
    path: str
    course: str | None
    learned_terms: list[str]
    score: float


@dataclass
class SummaryRecord:
    path: str
    course: str | None
    title: str
    summary: str
    key_points: list[str]
    important_terms: list[str]
    review_checklist: list[str]
    source: str


@dataclass
class QuestionRecord:
    course: str
    title: str
    question_type: str
    question: str
    answer: str
    explanation: str
    source_paths: list[str]
    source: str


def workspace_root(path: str | None) -> Path:
    return Path(path).expanduser().resolve() if path else Path.cwd().resolve()


def app_dir(root: Path) -> Path:
    return root / APP_DIR


def config_path(root: Path) -> Path:
    return app_dir(root) / CONFIG_FILE


def summaries_path(root: Path) -> Path:
    return app_dir(root) / SUMMARIES_FILE


def questions_path(root: Path) -> Path:
    return app_dir(root) / QUESTIONS_FILE


def migrate_existing_files(
    root: Path, courses: list[Course], records: list[MaterialRecord]
) -> bool:
    modified = False

    # 1. records に登録されているファイルを処理
    for record in records:
        path = material_path(root, record.path)
        if not path.exists() or not path.is_file():
            continue
        
        suffix = path.suffix.lower().lstrip(".")
        if not suffix:
            suffix = "file"
            
        parent_name = path.parent.name.lower()
        if parent_name == suffix:
            continue
            
        dest_dir = path.parent / suffix
        dest_dir.mkdir(parents=True, exist_ok=True)
        dest_path = unique_destination(dest_dir / path.name)
        
        try:
            shutil.move(str(path), str(dest_path))
            record.path = normalize_record_path(relative_to_root(root, dest_path))
            modified = True
        except Exception as e:
            print(f"Failed to migrate {path.name}: {e}", file=sys.stderr)

    # 2. データベースに載っていない物理ファイルも走査して移動
    target_dirs = [root / "その他"] + [root / c.folder for c in courses]
    for target_dir in target_dirs:
        if not target_dir.exists():
            continue
        try:
            for path in list(target_dir.iterdir()):
                if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
                    suffix = path.suffix.lower().lstrip(".") or "file"
                    dest_dir = target_dir / suffix
                    dest_dir.mkdir(parents=True, exist_ok=True)
                    dest_path = unique_destination(dest_dir / path.name)
                    try:
                        shutil.move(str(path), str(dest_path))
                        modified = True
                    except Exception as e:
                        print(f"Failed to migrate physical file {path.name}: {e}", file=sys.stderr)
        except Exception as e:
            print(f"Failed to read target dir {target_dir.name}: {e}", file=sys.stderr)

    return modified


def load_state(root: Path) -> tuple[list[Course], list[MaterialRecord]]:
    path = config_path(root)
    if not path.exists():
        return [], []
    data = json.loads(path.read_text(encoding="utf-8"))
    courses = [
        Course(
            name=item["name"],
            teacher=item.get("teacher", ""),
            keywords=item.get("keywords", []),
            folder=item["folder"],
            learned_terms=item.get("learned_terms", []),
            attendance=[
                AttendanceRecord(
                    date=att.get("date", ""),
                    class_round=att.get("class_round", 1),
                    status=att.get("status", "出席"),
                    memo=att.get("memo", ""),
                )
                for att in item.get("attendance", [])
            ],
            links=[
                LinkRecord(
                    title=lk.get("title", ""),
                    url=lk.get("url", ""),
                    memo=lk.get("memo", ""),
                )
                for lk in item.get("links", [])
            ],
        )
        for item in data.get("courses", [])
    ]
    records = [
        MaterialRecord(
            path=item["path"],
            course=item.get("course"),
            learned_terms=item.get("learned_terms", []),
            score=item.get("score", 0.0),
        )
        for item in data.get("materials", [])
    ]

    if migrate_existing_files(root, courses, records):
        save_state(root, courses, records)

    return courses, records


def load_courses(root: Path) -> list[Course]:
    courses, _records = load_state(root)
    return courses


def save_state(
    root: Path, courses: list[Course], records: list[MaterialRecord] | None = None
) -> None:
    app_dir(root).mkdir(parents=True, exist_ok=True)
    courses_payload = []
    for course in courses:
        c_dict = {
            "name": course.name,
            "teacher": course.teacher,
            "keywords": course.keywords,
            "folder": course.folder,
            "learned_terms": course.learned_terms,
            "attendance": [
                {
                    "date": att.date,
                    "class_round": att.class_round,
                    "status": att.status,
                    "memo": att.memo,
                }
                for att in course.attendance
            ],
            "links": [
                {
                    "title": lk.title,
                    "url": lk.url,
                    "memo": lk.memo,
                }
                for lk in course.links
            ],
        }
        courses_payload.append(c_dict)

    global_links_payload = []
    if config_path(root).exists():
        try:
            old_data = json.loads(config_path(root).read_text(encoding="utf-8"))
            global_links_payload = old_data.get("global_links", [])
        except Exception:
            pass

    payload = {
        "courses": courses_payload,
        "materials": [record.__dict__ for record in records or []],
        "global_links": global_links_payload,
    }
    config_path(root).write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def load_global_links(root: Path) -> list[LinkRecord]:
    path = config_path(root)
    if not path.exists():
        return []
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        raw_links = data.get("global_links", [])
        return [
            LinkRecord(
                title=item.get("title", ""),
                url=item.get("url", ""),
                memo=item.get("memo", "")
            )
            for item in raw_links
        ]
    except Exception:
        return []


def save_global_links(root: Path, links: list[LinkRecord]) -> None:
    path = config_path(root)
    if not path.exists():
        payload = {"courses": [], "materials": [], "global_links": []}
    else:
        try:
            payload = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            payload = {"courses": [], "materials": []}
            
    payload["global_links"] = [
        {
            "title": lk.title,
            "url": lk.url,
            "memo": lk.memo
        }
        for lk in links
    ]
    
    app_dir(root).mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def save_courses(root: Path, courses: list[Course]) -> None:
    _courses, records = load_state(root)
    save_state(root, courses, records)


def load_summaries(root: Path) -> list[SummaryRecord]:
    path = summaries_path(root)
    if not path.exists():
        return []
    data = json.loads(path.read_text(encoding="utf-8"))
    return [
        SummaryRecord(
            path=item["path"],
            course=item.get("course"),
            title=item.get("title", ""),
            summary=item.get("summary", ""),
            key_points=item.get("key_points", []),
            important_terms=item.get("important_terms", []),
            review_checklist=item.get("review_checklist", []),
            source=item.get("source", "unknown"),
        )
        for item in data.get("summaries", [])
    ]


def save_summaries(root: Path, summaries: list[SummaryRecord]) -> None:
    app_dir(root).mkdir(parents=True, exist_ok=True)
    payload = {"summaries": [summary.__dict__ for summary in summaries]}
    summaries_path(root).write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def load_questions(root: Path) -> list[QuestionRecord]:
    path = questions_path(root)
    if not path.exists():
        return []
    data = json.loads(path.read_text(encoding="utf-8"))
    return [
        QuestionRecord(
            course=item["course"],
            title=item.get("title", ""),
            question_type=item.get("question_type", "short_answer"),
            question=item.get("question", ""),
            answer=item.get("answer", ""),
            explanation=item.get("explanation", ""),
            source_paths=item.get("source_paths", []),
            source=item.get("source", "unknown"),
        )
        for item in data.get("questions", [])
    ]


def save_questions(root: Path, questions: list[QuestionRecord]) -> None:
    app_dir(root).mkdir(parents=True, exist_ok=True)
    payload = {"questions": [question.__dict__ for question in questions]}
    questions_path(root).write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def safe_folder_name(name: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", name).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned or "untitled_course"


def split_keywords(raw: str) -> list[str]:
    return [item.strip() for item in re.split(r"[,、\n]", raw) if item.strip()]


def add_course(args: argparse.Namespace) -> int:
    root = workspace_root(args.root)
    courses, records = load_state(root)
    if any(course.name == args.name for course in courses):
        print(f"授業 '{args.name}' はすでに登録されています。", file=sys.stderr)
        return 1

    folder = root / safe_folder_name(args.name)
    folder.mkdir(parents=True, exist_ok=True)
    course = Course(
        name=args.name,
        teacher=args.teacher or "",
        keywords=split_keywords(args.keywords or ""),
        folder=str(folder.relative_to(root)),
        learned_terms=[],
    )
    courses.append(course)
    save_state(root, courses, records)
    print(f"登録しました: {course.name} -> {folder}")
    return 0


def update_course(args: argparse.Namespace) -> int:
    root = workspace_root(args.root)
    courses, records = load_state(root)
    target_course = None
    for course in courses:
        if course.name == args.name:
            target_course = course
            break

    if not target_course:
        print(f"授業 '{args.name}' は登録されていません。", file=sys.stderr)
        return 1

    updated = False
    if getattr(args, "keywords", None) is not None:
        target_course.keywords = split_keywords(args.keywords)
        updated = True
    if getattr(args, "teacher", None) is not None:
        target_course.teacher = args.teacher
        updated = True

    if updated:
        save_state(root, courses, records)
        print(f"更新しました: {target_course.name}")
    else:
        print("変更する項目が指定されていません。")
    return 0





def list_courses(args: argparse.Namespace) -> int:
    root = workspace_root(args.root)
    courses = load_courses(root)
    if not courses:
        print("授業はまだ登録されていません。")
        return 0
    for course in courses:
        terms = ", ".join(course.keywords + course.learned_terms[:8])
        print(f"- {course.name} ({course.teacher})")
        print(f"  folder: {course.folder}")
        print(f"  terms: {terms or '(未設定)'}")
    return 0


def analyze_materials(args: argparse.Namespace) -> int:
    root = workspace_root(args.root)
    courses, records = load_state(root)
    if not courses:
        print("先に add-course で授業を登録してください。", file=sys.stderr)
        return 1

    summaries = load_summaries(root)
    targets = selected_materials(root, courses, records, args.course)
    if not targets:
        print("分析対象の資料が見つかりません。", file=sys.stderr)
        return 1

    count = 0
    for material, course in targets:
        if not args.force and find_summary(summaries, root, material):
            continue
        text = extract_text(material, args.ocr_language)
        if not text.strip():
            print(f"skip: テキストを抽出できませんでした: {material.name}")
            continue
        api_prov = getattr(args, "api_provider", "openai")
        summary = create_summary(
            root, material, course, text, args.model, args.local_only, api_provider=api_prov
        )
        upsert_summary(summaries, summary)
        try:
            write_summary_pdf(root, material, summary)
        except Exception as e:
            print(f"要約PDFの作成に失敗しましたが、要約データは保存します: {material.name}: {e}", file=sys.stderr)
        count += 1
        print(f"analyzed: {material.name} -> {summary.source}")

    save_summaries(root, summaries)
    print(f"完了: {count} 件の資料を分析しました。")
    try:
        if not hasattr(args, "output") or not args.output:
            args.output = "_site"
        build_site(args)
    except Exception as e:
        print(f"Failed to rebuild site after analyzing materials: {e}", file=sys.stderr)
    return 0


def generate_practice(args: argparse.Namespace) -> int:
    root = workspace_root(args.root)
    courses, records = load_state(root)
    if not courses:
        print("先に add-course で授業を登録してください。", file=sys.stderr)
        return 1

    summaries = load_summaries(root)
    questions = load_questions(root)
    target_courses = [find_course(courses, args.course)] if args.course else courses
    target_courses = [course for course in target_courses if course is not None]
    if not target_courses:
        print("対象の授業が見つかりません。", file=sys.stderr)
        return 1

    generated = 0
    if args.force:
        course_names = {course.name for course in target_courses}
        questions = [question for question in questions if question.course not in course_names]

    for course in target_courses:
        context, source_paths = practice_context(root, course, records, summaries)
        if not context.strip():
            print(f"skip: 問題生成に使える資料がありません: {course.name}")
            continue
        api_prov = getattr(args, "api_provider", "openai")
        new_questions = create_questions(
            course,
            context,
            source_paths,
            args.count,
            args.model,
            args.local_only,
            api_provider=api_prov,
        )
        questions.extend(new_questions)
        generated += len(new_questions)
        write_practice_markdown(root, course, new_questions)
        print(f"generated: {course.name} -> {len(new_questions)} 問")

    save_questions(root, questions)
    print(f"完了: {generated} 問を生成しました。")
    try:
        if not hasattr(args, "output") or not args.output:
            args.output = "_site"
        build_site(args)
    except Exception as e:
        print(f"Failed to rebuild site after generating practice: {e}", file=sys.stderr)
    return 0


def read_docx(path: Path) -> str:
    try:
        import docx
        doc = docx.Document(path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        print(f"docx {path.name} の読み取り中にエラーが発生しました: {e}", file=sys.stderr)
        return ""


def extract_text(path: Path, ocr_language: str) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".csv"}:
        return read_text_file(path)
    if suffix == ".pdf":
        return read_pdf(path, ocr_language)
    if suffix in {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff"}:
        return read_image(path, ocr_language)
    if suffix == ".docx":
        return read_docx(path)
    return ""


def read_text_file(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp932"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="ignore")


def read_pdf(path: Path, ocr_language: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        print(
            f"PDF読み取りには pypdf が必要です: {path.name}",
            file=sys.stderr,
        )
        return ""

    try:
        parts: list[str] = []
        reader = PdfReader(str(path))
        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)
                continue

            ocr_text = read_pdf_page_with_ocr(path, page_number, ocr_language)
            if ocr_text.strip():
                parts.append(ocr_text)
            else:
                parts.append("")
        return "\n".join(parts)
    except Exception as e:
        print(
            f"PDF {path.name} の読み取り中にエラーが発生しました: {e}",
            file=sys.stderr,
        )
        return ""


def read_pdf_page_with_ocr(path: Path, page_number: int, language: str) -> str:
    # Try pypdfium2 first
    try:
        import pypdfium2 as pdfium
        from PIL import Image, ImageOps
        import pytesseract

        user_tesseract = Path("C:/Users/kaiy2/OCR/tesseract.exe")
        if user_tesseract.exists():
            pytesseract.pytesseract.tesseract_cmd = str(user_tesseract)

        doc = pdfium.PdfDocument(str(path.resolve()))
        if 0 < page_number <= len(doc):
            page = doc[page_number - 1]
            bitmap = page.render(scale=2.0)
            pil_img = bitmap.to_pil()
            pil_img = ImageOps.grayscale(pil_img)
            pil_img = ImageOps.autocontrast(pil_img)
            ocr_text = pytesseract.image_to_string(pil_img, lang=language)
            if ocr_text.strip():
                return ocr_text
    except Exception as e:
        pass

    # Fallback to pdftoppm
    pdftoppm = find_pdftoppm()
    if not pdftoppm:
        print(
            f"PDF {path.name} のページ {page_number} は画像PDFの可能性があります。"
            "OCRには pypdfium2 または Poppler の pdftoppm が必要です。",
            file=sys.stderr,
        )
        return ""

    with tempfile.TemporaryDirectory(prefix="class_materials_pdf_ocr_") as temp_dir:
        output_prefix = Path(temp_dir) / "page"
        command = [
            pdftoppm,
            "-f",
            str(page_number),
            "-l",
            str(page_number),
            "-r",
            "300",
            "-png",
            str(path),
            str(output_prefix),
        ]
        try:
            subprocess.run(
                command,
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.PIPE,
                text=True,
            )
        except subprocess.CalledProcessError as e:
            detail = (e.stderr or "").strip()
            print(
                f"PDF {path.name} のページ {page_number} をOCR用画像に変換できませんでした: "
                f"{detail or e}",
                file=sys.stderr,
            )
            return ""
        except Exception as e:
            print(
                f"PDF {path.name} のページ {page_number} をOCR用画像に変換できませんでした: {e}",
                file=sys.stderr,
            )
            return ""

        images = sorted(Path(temp_dir).glob("page-*.png"))
        if not images:
            print(
                f"PDF {path.name} のページ {page_number} からOCR用画像を生成できませんでした。",
                file=sys.stderr,
            )
            return ""
        return read_image(images[0], language)


def find_pdftoppm() -> str | None:
    discovered = shutil.which("pdftoppm")
    candidates: list[Path] = []
    if discovered:
        discovered_path = Path(discovered)
        runtime_root = discovered_path.parent.parent
        candidates.append(runtime_root / "native" / "poppler" / "Library" / "bin" / "pdftoppm.exe")
        candidates.append(discovered_path)

    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return discovered


def read_image(path: Path, language: str) -> str:
    try:
        import pytesseract
        from PIL import Image, ImageOps
        
        user_tesseract = Path("C:/Users/kaiy2/OCR/tesseract.exe")
        if user_tesseract.exists():
            pytesseract.pytesseract.tesseract_cmd = str(user_tesseract)
    except ImportError:
        print(
            f"画像OCRには pillow と pytesseract が必要です: {path.name}",
            file=sys.stderr,
        )
        return ""

    try:
        image = Image.open(path)
        image = ImageOps.grayscale(image)
        image = ImageOps.autocontrast(image)
        return pytesseract.image_to_string(image, lang=language)
    except Exception as e:
        print(
            f"画像 {path.name} のOCR処理中にエラーが発生しました: {e}",
            file=sys.stderr,
        )
        return ""


VALID_3_LETTER_WORDS = {
    "git", "web", "cpu", "ram", "gpu", "os", "pdf", "api", "sql", "url", "csv", "xml", "txt", "app", 
    "dev", "run", "key", "use", "get", "set", "add", "new", "map", "zip", "net", "dns", "tcp", "udp", 
    "ssl", "tls", "ssh", "ftp", "mac", "vpn", "ips", "ids", "lan", "wan", "iot", "dom", "css", "cpp",
    "bin", "hex", "oct", "dec", "bit", "int", "var", "val", "arg", "fun", "def", "cls", "lib", "api",
    "cmd", "dir", "log", "err", "out", "src", "sys", "usr", "doc", "cfg", "env", "md5", "sha", "aes",
    "rsa", "des", "xls", "ppt", "png", "jpg", "gif", "bmp", "svg", "wav", "mp3", "mp4", "avi", "mov", 
    "mkv", "tar", "rar", "7z"
}

def tokenize(text: str) -> list[str]:
    normalized = text.lower()
    tokens = re.findall(r"[a-zA-Z][a-zA-Z0-9_-]{2,}|[一-龥ぁ-んァ-ンー]{2,}", normalized)
    result = []
    for token in tokens:
        if token in JAPANESE_STOPWORDS or token in ENGLISH_STOPWORDS:
            continue
        if re.search(r"\d", token):
            continue
        if token.startswith(("-", "_")) or token.endswith(("-", "_")):
            continue
        if re.match(r"^[ぁ-んー]+$", token) and len(token) <= 3:
            continue
        if re.match(r"^[a-zA-Z_-]+$", token):
            pure_alpha = token.replace("-", "").replace("_", "")
            if not re.search(r"[aeiouy]", pure_alpha):
                continue
            if len(pure_alpha) == 3 and pure_alpha not in VALID_3_LETTER_WORDS:
                continue
        result.append(token)
    return result


def best_course(text: str, courses: list[Course]) -> tuple[Course | None, float, list[str]]:
    document_terms = tokenize(text)
    document_counts = Counter(document_terms)
    if not document_counts:
        return None, 0.0, []

    best: Course | None = None
    best_score = 0.0
    for course in courses:
        profile_terms = tokenize(
            " ".join([course.name, course.teacher, *course.keywords, *course.learned_terms])
        )
        profile_counts = Counter(profile_terms)
        score = cosine_similarity(document_counts, profile_counts)
        keyword_bonus = direct_keyword_matches(text, course)
        score += keyword_bonus * 0.15
        if score > best_score:
            best = course
            best_score = score

    learned = [term for term, _ in document_counts.most_common(12)]
    return best, best_score, learned


def direct_keyword_matches(text: str, course: Course) -> int:
    lowered = text.lower()
    terms = [course.name, course.teacher, *course.keywords, *course.learned_terms]
    return sum(1 for term in terms if term and term.lower() in lowered)


def cosine_similarity(left: Counter[str], right: Counter[str]) -> float:
    if not left or not right:
        return 0.0
    shared = set(left) & set(right)
    numerator = sum(left[term] * right[term] for term in shared)
    left_norm = math.sqrt(sum(value * value for value in left.values()))
    right_norm = math.sqrt(sum(value * value for value in right.values()))
    if left_norm == 0 or right_norm == 0:
        return 0.0
    return numerator / (left_norm * right_norm)


def iter_materials(paths: Iterable[str], recursive: bool) -> Iterable[Path]:
    for raw in paths:
        path = Path(raw).expanduser().resolve()
        if path.is_file():
            yield path
            continue
        if path.is_dir():
            pattern = "**/*" if recursive else "*"
            for child in path.glob(pattern):
                if child.is_file() and child.suffix.lower() in SUPPORTED_EXTENSIONS:
                    yield child


def sort_materials(args: argparse.Namespace) -> int:
    root = workspace_root(args.root)
    courses, records = load_state(root)
    if not courses:
        print("先に add-course で授業を登録してください。", file=sys.stderr)
        return 1

    unknown_dir = root / "その他"
    moved = 0
    for material in iter_materials(args.inputs, args.recursive):
        text = extract_text(material, args.ocr_language)
        course, score, learned_terms = best_course(text, courses)
        suffix_folder = material.suffix.lower().lstrip(".")
        if not suffix_folder:
            suffix_folder = "file"

        if course is None or score < args.threshold:
            destination_dir = unknown_dir / suffix_folder
            label = "その他"
        else:
            destination_dir = root / course.folder / suffix_folder
            label = course.name

        destination_dir.mkdir(parents=True, exist_ok=True)
        destination = unique_destination(destination_dir / material.name)
        if args.copy:
            shutil.copy2(material, destination)
            action = "copied"
        else:
            shutil.move(str(material), str(destination))
            action = "moved"
        
        course_name = course.name if course is not None and score >= args.threshold else None
        upsert_material_record(
            records,
            root,
            destination,
            course_name,
            learned_terms,
            score,
        )
        if course_name and course:
            recalculate_course_learned_terms(course, records)
        moved += 1
        print(f"{action}: {material.name} -> {label} (score={score:.3f})")

    save_state(root, courses, records)
    print(f"完了: {moved} 件の資料を処理しました。")
    return 0


def reassign_material(args: argparse.Namespace) -> int:
    root = workspace_root(args.root)
    courses, records = load_state(root)

    material = Path(args.material).expanduser().resolve()
    if not material.exists() or not material.is_file():
        print(f"資料ファイルが見つかりません: {material}", file=sys.stderr)
        return 1

    suffix_folder = material.suffix.lower().lstrip(".")
    if not suffix_folder:
        suffix_folder = "file"

    is_to_other = (args.to == "その他" or args.to == "_未分類" or args.to == "未分類")
    if is_to_other:
        target = None
        destination_dir = root / "その他" / suffix_folder
    else:
        target = find_course(courses, args.to)
        if target is None:
             print(f"移動先の授業が見つかりません: {args.to}", file=sys.stderr)
             return 1
        destination_dir = root / target.folder / suffix_folder

    record = find_material_record(records, root, material)
    source_course_name = record.course if record else infer_course_from_path(root, courses, material)
    source = find_course(courses, source_course_name) if source_course_name else None
    
    if (source is None and target is None) or (source and target and source.name == target.name):
        print("資料はすでに指定された場所にあります。")
        return 0

    if record:
        learned_terms = record.learned_terms
        score = record.score
    else:
        text = extract_text(material, args.ocr_language)
        learned_terms = [term for term, _count in Counter(tokenize(text)).most_common(12)]
        score = 0.0

    destination_dir.mkdir(parents=True, exist_ok=True)
    destination = unique_destination(destination_dir / material.name)
    shutil.move(str(material), str(destination))

    target_name = target.name if target else None

    if record:
        records[:] = [
            item
            for item in records
            if normalize_record_path(item.path) != normalize_record_path(record.path)
        ]
    upsert_material_record(records, root, destination, target_name, learned_terms, score)

    if source:
        recalculate_course_learned_terms(source, records)
    if target:
        recalculate_course_learned_terms(target, records)

    save_state(root, courses, records)
    source_label = source.name if source else "その他"
    target_label = target.name if target else "その他"
    print(f"変更しました: {material.name} ({source_label} -> {target_label})")
    return 0


def delete_material(args: argparse.Namespace) -> int:
    root = workspace_root(args.root)
    courses, records = load_state(root)

    material = Path(args.material).expanduser().resolve()
    if not material.exists() or not material.is_file():
        print(f"資料ファイルが見つかりません: {material}", file=sys.stderr)
        return 1

    record = find_material_record(records, root, material)
    source_course_name = record.course if record else infer_course_from_path(root, courses, material)
    source = find_course(courses, source_course_name) if source_course_name else None

    if record:
        learned_terms = record.learned_terms
    else:
        text = extract_text(material, args.ocr_language)
        learned_terms = [term for term, _count in Counter(tokenize(text)).most_common(12)]

    # Remove the physical file
    material.unlink()

    # Remove the record
    if record:
        records[:] = [
            item
            for item in records
            if normalize_record_path(item.path) != normalize_record_path(record.path)
        ]

    # Recalculate learned terms for the source course
    if source:
        recalculate_course_learned_terms(source, records)

    save_state(root, courses, records)
    source_label = source.name if source else "その他"
    print(f"削除しました: {material.name} (元所属: {source_label})")
    return 0


def clear_analysis(args: argparse.Namespace) -> int:
    root = workspace_root(args.root)
    courses, records = load_state(root)
    
    target_course_name = getattr(args, "course", None)
    target_material_path_str = getattr(args, "material", None)
    
    # 1. 要約のクリア
    summaries = load_summaries(root)
    remaining_summaries = []
    cleared_summaries_count = 0
    
    for summary in summaries:
        should_clear = False
        if target_course_name and summary.course == target_course_name:
            should_clear = True
        elif target_material_path_str:
            mat_path = Path(target_material_path_str).expanduser().resolve()
            record_path = material_path(root, summary.path).resolve()
            if mat_path == record_path:
                should_clear = True
        elif not target_course_name and not target_material_path_str:
            should_clear = True
            
        if should_clear:
            path = material_path(root, summary.path)
            md_path = path.with_suffix(path.suffix + ".summary.md")
            if md_path.exists():
                try:
                    md_path.unlink()
                except Exception as e:
                    print(f"Failed to delete markdown file {md_path}: {e}", file=sys.stderr)
            pdf_path = path.with_suffix(path.suffix + ".summary.pdf")
            if pdf_path.exists():
                try:
                    pdf_path.unlink()
                except Exception as e:
                    print(f"Failed to delete PDF file {pdf_path}: {e}", file=sys.stderr)
            cleared_summaries_count += 1
        else:
            remaining_summaries.append(summary)
            
    save_summaries(root, remaining_summaries)
    print(f"要約を {cleared_summaries_count} 件削除しました。")
    
    # 2. 練習問題のクリア
    questions = load_questions(root)
    remaining_questions = []
    cleared_questions_count = 0
    
    target_courses = []
    if target_course_name:
        course = find_course(courses, target_course_name)
        if course:
            target_courses.append(course)
    elif not target_material_path_str:
        target_courses = courses
        
    for question in questions:
        should_clear = False
        if target_course_name and question.course == target_course_name:
            should_clear = True
        elif not target_course_name and not target_material_path_str:
            should_clear = True
            
        if should_clear:
            cleared_questions_count += 1
        else:
            remaining_questions.append(question)
            
    for course in target_courses:
        pq_path = root / course.folder / "practice_questions.md"
        if pq_path.exists():
            try:
                pq_path.unlink()
            except Exception as e:
                print(f"Failed to delete practice questions markdown {pq_path}: {e}", file=sys.stderr)
                
    save_questions(root, remaining_questions)
    print(f"練習問題を {cleared_questions_count} 問削除しました。")

    try:
        site_args = argparse.Namespace(
            root=getattr(args, "root", None),
            output=getattr(args, "output", "_site")
        )
        build_site(site_args)
    except Exception as e:
        print(f"Failed to rebuild site after clearing analysis: {e}", file=sys.stderr)
    
    return 0


def build_site(args: argparse.Namespace) -> int:
    root = workspace_root(args.root)
    courses, records = load_state(root)
    if not courses:
        print("先に add-course で授業を登録してください。", file=sys.stderr)
        return 1

    output_dir = (root / args.output).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    index_path = output_dir / "index.html"
    summaries = load_summaries(root)
    questions = load_questions(root)
    index_path.write_text(
        render_site(root, output_dir, courses, records, summaries, questions),
        encoding="utf-8",
    )
    print(f"作成しました: {index_path}")
    return 0


def render_site(
    root: Path,
    output_dir: Path,
    courses: list[Course],
    records: list[MaterialRecord],
    summaries: list[SummaryRecord],
    questions: list[QuestionRecord],
) -> str:
    course_sections = "\n".join(
        render_course_section(root, output_dir, course, records, summaries, questions)
        for course in courses
    )
    unknown_section = render_unknown_section(root, output_dir, records, summaries)
    total_materials = sum(len(course_materials(root, course, records)) for course in courses)
    total_materials += len(unknown_materials(root, records))
    course_nav = "\n".join(
        f'<a href="#{html_id(course.name)}">{escape(course.name)}</a>' for course in courses
    )

    global_links = load_global_links(root)
    global_links_html = ""
    if global_links:
        link_items = []
        for lk in global_links:
            memo_html = f'<span style="font-size:0.85rem; color:var(--muted); margin-left: 10px;">{escape(lk.memo)}</span>' if lk.memo else ""
            link_items.append(f"""
            <li style="margin-bottom: 8px; list-style: none; display: flex; align-items: center; gap: 8px;">
              <a href="{escape(lk.url)}" target="_blank" style="color:var(--accent); font-weight:bold; text-decoration:none; display: inline-flex; align-items: center; gap: 4px;">
                <span>🔗</span> <span style="text-decoration: underline;">{escape(lk.title)}</span>
              </a>
              {f'<span style="color:var(--line)">|</span>' if memo_html else ""}
              {memo_html}
            </li>
            """)
        link_items_html = "\n".join(link_items)
        global_links_html = f"""
        <section class="global-links-section" style="background: var(--surface); border: 1px solid var(--line); border-radius: 8px; padding: 24px; margin-bottom: 24px;">
          <h2 style="margin: 0 0 16px 0; font-size: 1.25rem; color: var(--accent); display: flex; align-items: center; gap: 8px;">🔗 共通参考URLリンク集</h2>
          <ul style="margin: 0; padding: 0;">
            {link_items_html}
          </ul>
        </section>
        """

    return f"""<!doctype html>
<html lang="ja">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>授業資料ライブラリ</title>
  <style>
    :root {{
      color-scheme: light;
      --ink: #17202a;
      --muted: #627083;
      --line: #d8dee8;
      --surface: #ffffff;
      --soft: #f5f7fa;
      --accent: #0f766e;
      --accent-soft: #d8f3ef;
      --warn: #9a5b00;
      --warn-soft: #fff3d8;
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      color: var(--ink);
      background: var(--soft);
      font-family: "Yu Gothic UI", "Meiryo", system-ui, sans-serif;
      line-height: 1.5;
    }}
    header {{
      background: var(--surface);
      border-bottom: 1px solid var(--line);
      padding: 28px min(5vw, 48px) 20px;
      position: sticky;
      top: 0;
      z-index: 2;
    }}
    h1 {{
      font-size: clamp(1.5rem, 2vw, 2rem);
      margin: 0 0 14px;
      letter-spacing: 0;
    }}
    .toolbar {{
      align-items: center;
      display: grid;
      gap: 12px;
      grid-template-columns: minmax(220px, 420px) 1fr auto;
    }}
    input[type="search"] {{
      border: 1px solid var(--line);
      border-radius: 6px;
      font: inherit;
      min-height: 42px;
      padding: 0 12px;
      width: 100%;
    }}
    nav {{
      display: flex;
      flex-wrap: wrap;
      gap: 8px;
      min-width: 0;
    }}
    nav a, .stat {{
      background: var(--soft);
      border: 1px solid var(--line);
      border-radius: 6px;
      color: var(--ink);
      padding: 8px 10px;
      text-decoration: none;
      white-space: nowrap;
    }}
    .status-badge {{
      display: inline-block;
      padding: 2px 6px;
      font-size: 0.8rem;
      border-radius: 4px;
      font-weight: 500;
      line-height: 1.2;
    }}
    .status-badge.present {{
      background: #d8f3ef;
      color: #0f766e;
    }}
    .status-badge.late {{
      background: #fef3c7;
      color: #b25e00;
    }}
    .status-badge.absent {{
      background: #fee2e2;
      color: #b91c1c;
    }}
    .status-badge.excused {{
      background: #dbeafe;
      color: #1d4ed8;
    }}
    .attendance-table th {{
      border-bottom: 2px solid var(--line);
    }}
    .attendance-table td {{
      padding: 8px 12px;
      border-bottom: 1px solid var(--line);
    }}
    .attendance-table tr:last-child td {{
      border-bottom: none;
    }}
    main {{
      display: grid;
      gap: 22px;
      padding: 24px min(5vw, 48px) 48px;
    }}
    section {{
      background: var(--surface);
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 20px;
    }}
    .course-head {{
      align-items: start;
      display: flex;
      gap: 16px;
      justify-content: space-between;
      margin-bottom: 16px;
    }}
    h2 {{
      font-size: 1.2rem;
      margin: 0 0 4px;
      letter-spacing: 0;
    }}
    .teacher, .meta {{
      color: var(--muted);
      font-size: 0.92rem;
    }}
    .chips {{
      display: flex;
      flex-wrap: wrap;
      gap: 6px;
      margin-top: 10px;
    }}
    .chip {{
      background: var(--accent-soft);
      border-radius: 999px;
      color: #07534d;
      font-size: 0.82rem;
      padding: 4px 8px;
    }}
    .grid {{
      display: grid;
      gap: 12px;
      grid-template-columns: repeat(auto-fill, minmax(240px, 1fr));
    }}
    .material {{
      border: 1px solid var(--line);
      border-radius: 8px;
      display: grid;
      gap: 10px;
      min-height: 128px;
      padding: 12px;
      text-decoration: none;
      color: inherit;
      background: #fff;
    }}
    .material:hover {{
      border-color: var(--accent);
      box-shadow: 0 4px 18px rgba(15, 118, 110, 0.12);
    }}
    .btn {{
      background: var(--accent-soft);
      border: 1px solid var(--line);
      border-radius: 6px;
      color: #07534d;
      padding: 6px 12px;
      text-decoration: none;
      font-size: 0.85rem;
      text-align: center;
      display: inline-block;
      cursor: pointer;
      font-weight: 500;
    }}
    .btn:hover {{
      background: var(--accent);
      color: white;
    }}
    .summary-details {{
      width: 100%;
      margin-top: 8px;
    }}
    .summary-details summary {{
      outline: none;
      list-style: none;
    }}
    .summary-details summary::-webkit-details-marker {{
      display: none;
    }}
    .summary-content {{
      border-top: 1px solid var(--line);
      margin-top: 10px;
      padding-top: 10px;
      font-size: 0.88rem;
    }}
    .summary-content h4 {{
      margin: 8px 0 4px;
      font-size: 0.92rem;
      color: var(--accent);
    }}
    .summary-content p {{
      margin: 0 0 8px;
      color: var(--ink);
    }}
    .summary-content ul {{
      margin: 0 0 8px;
      padding-left: 20px;
    }}
    .preview {{
      align-items: center;
      background: var(--soft);
      border-radius: 6px;
      display: flex;
      height: 76px;
      justify-content: center;
      overflow: hidden;
    }}
    .preview img {{
      height: 100%;
      object-fit: cover;
      width: 100%;
    }}
    .badge {{
      background: var(--ink);
      border-radius: 6px;
      color: white;
      font-size: 0.8rem;
      font-weight: 700;
      padding: 5px 8px;
      text-transform: uppercase;
    }}
    .name {{
      font-weight: 700;
      overflow-wrap: anywhere;
    }}
    .terms {{
      color: var(--muted);
      font-size: 0.86rem;
      overflow-wrap: anywhere;
    }}
    .empty {{
      border: 1px dashed var(--line);
      border-radius: 8px;
      color: var(--muted);
      padding: 16px;
    }}
    .insights {{
      display: grid;
      gap: 12px;
      grid-template-columns: repeat(auto-fit, minmax(260px, 1fr));
      margin-bottom: 16px;
    }}
    .panel {{
      background: var(--soft);
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 12px;
    }}
    .panel h3 {{
      font-size: 1rem;
      margin: 0 0 8px;
      letter-spacing: 0;
    }}
    .panel ul {{
      margin: 8px 0 0;
      padding-left: 20px;
    }}
    .question {{
      border-top: 1px solid var(--line);
      padding-top: 10px;
      margin-top: 10px;
    }}
    .unknown {{
      background: var(--warn-soft);
      border-color: #f0cf8b;
    }}
    .unknown .chip {{
      background: #ffe2a7;
      color: var(--warn);
    }}
    @media (max-width: 760px) {{
      header {{ position: static; }}
      .toolbar {{ grid-template-columns: 1fr; }}
      .course-head {{ display: grid; }}
    }}
  </style>
</head>
<body>
  <header>
    <h1>授業資料ライブラリ</h1>
    <div class="toolbar">
      <input id="search" type="search" placeholder="授業名・教員名・キーワード・資料名で検索">
      <nav>{course_nav}</nav>
      <div class="stat">{len(courses)} 授業 / {total_materials} 資料</div>
    </div>
  </header>
  <main style="padding: 24px min(5vw, 48px);">
    {global_links_html}
    {course_sections}
    {unknown_section}
  </main>
  <script>
    const search = document.getElementById("search");
    const items = Array.from(document.querySelectorAll("[data-search]"));
    search.addEventListener("input", () => {{
      const query = search.value.trim().toLowerCase();
      for (const item of items) {{
        item.hidden = query.length > 0 && !item.dataset.search.includes(query);
      }}
      for (const section of document.querySelectorAll("section")) {{
        const visible = Array.from(section.querySelectorAll(".material")).some((item) => !item.hidden);
        section.hidden = query.length > 0 && !visible && !section.dataset.search.includes(query);
      }}
    }});
  </script>
</body>
</html>
"""


def render_course_section(
    root: Path,
    output_dir: Path,
    course: Course,
    records: list[MaterialRecord],
    summaries: list[SummaryRecord],
    questions: list[QuestionRecord],
) -> str:
    materials = course_materials(root, course, records)
    cards = "\n".join(render_material_card(root, output_dir, item, summaries) for item in materials)
    if not cards:
        cards = '<div class="empty">資料はまだありません。</div>'
    terms = course.keywords + course.learned_terms
    chips = render_chips(terms[:16])
    insights = render_course_insights(course, summaries, questions)
    
    # 出席状況テーブルの生成
    attendance_html = ""
    if hasattr(course, "attendance") and course.attendance:
        total = len(course.attendance)
        presents = sum(1 for att in course.attendance if att.status == "出席")
        lates = sum(1 for att in course.attendance if att.status == "遅刻")
        absents = sum(1 for att in course.attendance if att.status == "欠席")
        excuseds = sum(1 for att in course.attendance if att.status == "公欠")
        rate = (presents + lates) / total * 100 if total > 0 else 0.0
        
        rows = []
        for att in course.attendance:
            status_class = "present"
            if att.status == "欠席":
                status_class = "absent"
            elif att.status == "遅刻":
                status_class = "late"
            elif att.status == "公欠":
                status_class = "excused"
            
            rows.append(f"""
            <tr>
              <td style="padding: 8px 12px; border-bottom: 1px solid var(--line);">第 {att.class_round} 回</td>
              <td style="padding: 8px 12px; border-bottom: 1px solid var(--line);">{escape(att.date)}</td>
              <td style="padding: 8px 12px; border-bottom: 1px solid var(--line);"><span class="status-badge {status_class}">{escape(att.status)}</span></td>
              <td style="padding: 8px 12px; border-bottom: 1px solid var(--line);">{escape(att.memo) if att.memo else '<span style="color:var(--muted)">-</span>'}</td>
            </tr>
            """)
        rows_html = "\n".join(rows)
        
        attendance_html = f"""
        <div class="attendance-section" style="margin-top: 15px; margin-bottom: 20px; background: var(--soft); border: 1px solid var(--line); border-radius: 6px; padding: 15px;">
          <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 10px; flex-wrap: wrap; gap: 10px;">
            <h3 style="margin: 0; font-size: 1.1rem; color: var(--accent);">📊 出席状況</h3>
            <div style="display: flex; gap: 12px; font-size: 0.9rem; color: var(--muted); flex-wrap: wrap;">
              <span>総回数: <strong>{total}</strong></span>
              <span>出席: <strong style="color:#0f766e">{presents}</strong></span>
              <span>遅刻: <strong style="color:#b25e00">{lates}</strong></span>
              <span>欠席: <strong style="color:#b91c1c">{absents}</strong></span>
              <span>公欠: <strong style="color:#1d4ed8">{excuseds}</strong></span>
              <span style="margin-left: 10px;">出席率: <strong style="color:var(--ink)">{rate:.1f}%</strong></span>
            </div>
          </div>
          <div style="width: 100%; height: 6px; background: #e2e8f0; border-radius: 3px; overflow: hidden; margin-bottom: 15px;">
            <div style="width: {rate}%; height: 100%; background: var(--accent); transition: width 0.3s ease;"></div>
          </div>
          <table class="attendance-table" style="width: 100%; border-collapse: collapse; font-size: 0.9rem; text-align: left; background: #fff; border-radius: 4px; overflow: hidden;">
            <thead>
              <tr style="border-bottom: 2px solid var(--line); background: var(--soft); color: var(--muted);">
                <th style="padding: 8px 12px; width: 80px;">回数</th>
                <th style="padding: 8px 12px; width: 120px;">日付</th>
                <th style="padding: 8px 12px; width: 80px;">状況</th>
                <th style="padding: 8px 12px;">授業メモ・課題など</th>
              </tr>
            </thead>
            <tbody>
              {rows_html}
            </tbody>
          </table>
        </div>
        """
        
    search_text = " ".join([course.name, course.teacher, *terms]).lower()
    return f"""<section id="{html_id(course.name)}" data-search="{escape_attr(search_text)}">
  <div class="course-head">
    <div>
      <h2>{escape(course.name)}</h2>
      <div class="teacher">{escape(course.teacher) if course.teacher else "教員名未登録"}</div>
      {chips}
    </div>
    <div class="meta">{len(materials)} 資料</div>
  </div>
  {attendance_html}
  {insights}
  <div class="grid">{cards}</div>
</section>"""


def render_unknown_section(
    root: Path, output_dir: Path, records: list[MaterialRecord], summaries: list[SummaryRecord]
) -> str:
    materials = unknown_materials(root, records)
    if not materials:
        return ""
    cards = "\n".join(render_material_card(root, output_dir, item, summaries) for item in materials)
    return f"""<section id="unknown" class="unknown" data-search="その他">
  <div class="course-head">
    <div>
      <h2>その他</h2>
      <div class="teacher">確認が必要な資料</div>
      {render_chips(["要確認"])}
    </div>
    <div class="meta">{len(materials)} 資料</div>
  </div>
  <div class="grid">{cards}</div>
</section>"""


def render_course_insights(
    course: Course, summaries: list[SummaryRecord], questions: list[QuestionRecord]
) -> str:
    course_summaries = [summary for summary in summaries if summary.course == course.name]
    course_questions = [question for question in questions if question.course == course.name]
    if not course_summaries and not course_questions:
        return ""

    summary_html = ""
    if course_summaries:
        summary_items = "\n".join(
            f"""<details class="summary-details">
  <summary><strong>{escape(summary.title)}</strong></summary>
  <div class="summary-content">
    <h4>要約</h4>
    <p>{escape(summary.summary)}</p>
    {render_text_list("要点", summary.key_points)}
    {render_text_list("重要語句", summary.important_terms)}
    {render_text_list("復習チェック", summary.review_checklist)}
  </div>
</details>"""
            for summary in course_summaries
        )
        summary_html = f"""<div class="panel">
  <h3>AI要約 ({len(course_summaries)}件)</h3>
  {summary_items}
</div>"""

    question_html = ""
    if course_questions:
        question_items = "\n".join(
            f"""<div class="question">
  <strong>{escape(question.title)}</strong>
  <div style="font-size: 0.8em; color: #6b7280; margin-bottom: 6px;">生成元: {escape(question.source)}</div>
  <div>{escape(question.question)}</div>
  <details><summary>解答と解説</summary><p>{escape(question.answer)}</p><p>{escape(question.explanation)}</p></details>
</div>"""
            for question in course_questions
        )
        question_html = f"""<div class="panel">
  <h3>練習問題 ({len(course_questions)}問)</h3>
  {question_items}
</div>"""

    return f'<div class="insights">{summary_html}{question_html}</div>'


def render_text_list(title: str, items: list[str]) -> str:
    if not items:
        return ""
    item_html = "".join(f"<li>{escape(item)}</li>" for item in items)
    return f"<h4>{escape(title)}</h4><ul>{item_html}</ul>"


def get_thumbnail_cache_path(root: Path, pdf_path: Path) -> Path:
    import hashlib
    rel_path = relative_to_root(root, pdf_path)
    path_hash = hashlib.md5(rel_path.encode("utf-8")).hexdigest()
    return root / ".class_materials" / "thumbnails" / f"{path_hash}.png"


def ensure_pdf_thumbnail(root: Path, pdf_path: Path) -> Path | None:
    cache_path = get_thumbnail_cache_path(root, pdf_path)
    if cache_path.exists():
        return cache_path

    cache_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        import pypdfium2 as pdfium
        from PIL import Image
    except ImportError as e:
        print(f"ImportError in ensure_pdf_thumbnail: {e}", file=sys.stderr)
        return None

    try:
        doc = pdfium.PdfDocument(str(pdf_path.resolve()))
        if len(doc) == 0:
            return None
        page = doc[0]
        bitmap = page.render(scale=1.5)
        pil_img = bitmap.to_pil()
        pil_img.thumbnail((250, 350))
        pil_img.save(cache_path, "PNG")
        return cache_path
    except Exception as e:
        print(f"PDFサムネイル生成エラー ({pdf_path.name}): {e}", file=sys.stderr)
        return None


def render_material_card(
    root: Path, output_dir: Path, material: Path, summaries: list[SummaryRecord]
) -> str:
    suffix = material.suffix.lower().lstrip(".") or "file"
    href = href_between(output_dir, material)
    preview = render_preview(root, output_dir, material, suffix)
    search_text = " ".join([material.name, suffix, str(material.parent.name)]).lower()

    summary = find_summary(summaries, root, material)
    summary_html = ""
    if summary:
        summary_text = f"""<details class="summary-details">
  <summary><span class="btn">要約を表示</span></summary>
  <div class="summary-content">
    <h4>要約</h4>
    <p>{escape(summary.summary)}</p>
    {render_text_list("要点", summary.key_points)}
    {render_text_list("重要語句", summary.important_terms)}
  </div>
</details>"""
        pdf_path = material.with_suffix(material.suffix + ".summary.pdf")
        pdf_link = ""
        if pdf_path.exists():
            pdf_href = href_between(output_dir, pdf_path)
            pdf_link = f'<a class="btn" href="{escape_attr(pdf_href)}" target="_blank">要約PDFを開く</a>'
        summary_html = f"{pdf_link}{summary_text}"

    return f"""<div class="material" data-search="{escape_attr(search_text)}">
  <div style="display: flex; gap: 10px; width: 100%; align-items: start;">
    {preview}
    <div style="flex: 1; min-width: 0;">
      <div class="name">{escape(material.name)}</div>
      <div class="terms">{escape(relative_to_root(root, material))}</div>
    </div>
  </div>
  <div style="display: flex; flex-wrap: wrap; gap: 8px; width: 100%; margin-top: auto; padding-top: 8px;">
    <a class="btn" href="{escape_attr(href)}" target="_blank">資料を開く</a>
    {summary_html}
  </div>
</div>"""


def render_preview(root: Path, output_dir: Path, material: Path, suffix: str) -> str:
    if suffix in {"png", "jpg", "jpeg", "bmp", "tif", "tiff"}:
        src = href_between(output_dir, material)
        return f'<div class="preview"><img src="{escape_attr(src)}" alt=""></div>'
    
    if suffix == "pdf":
        thumb_cache = ensure_pdf_thumbnail(root, material)
        if thumb_cache and thumb_cache.exists():
            thumb_dest_dir = output_dir / "_thumbnails"
            thumb_dest_dir.mkdir(parents=True, exist_ok=True)
            thumb_dest = thumb_dest_dir / thumb_cache.name
            if not thumb_dest.exists():
                shutil.copy2(thumb_cache, thumb_dest)
            
            src = href_between(output_dir, thumb_dest)
            return f'<div class="preview"><img src="{escape_attr(src)}" alt=""></div>'

    return f'<div class="preview"><span class="badge">{escape(suffix)}</span></div>'


def render_chips(terms: list[str]) -> str:
    if not terms:
        return ""
    chips = "\n".join(f'<span class="chip">{escape(term)}</span>' for term in unique_terms(terms))
    return f'<div class="chips">{chips}</div>'


def course_materials(
    root: Path, course: Course, records: list[MaterialRecord]
) -> list[Path]:
    course_dir = root / course.folder
    materials: list[Path] = []
    for record in records:
        if record.course == course.name:
            path = material_path(root, record.path)
            if path.exists():
                materials.append(path)
    if course_dir.exists():
        known = {normalize_record_path(relative_to_root(root, item)) for item in materials}
        for path in sorted(course_dir.glob("**/*"), key=lambda item: item.name.lower()):
            if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
                if path.name.lower().endswith(".summary.pdf"):
                    continue
                relative = normalize_record_path(relative_to_root(root, path))
                if relative not in known:
                    materials.append(path)
    return sorted(materials, key=lambda item: item.name.lower())


def unknown_materials(root: Path, records: list[MaterialRecord]) -> list[Path]:
    materials: list[Path] = []
    for record in records:
        if record.course is None:
            path = material_path(root, record.path)
            if path.exists():
                materials.append(path)
    unknown_dir = root / "その他"
    if unknown_dir.exists():
        known = {normalize_record_path(relative_to_root(root, item)) for item in materials}
        for path in sorted(unknown_dir.glob("**/*"), key=lambda item: item.name.lower()):
            if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
                relative = normalize_record_path(relative_to_root(root, path))
                if relative not in known:
                    materials.append(path)
    return sorted(materials, key=lambda item: item.name.lower())


def material_path(root: Path, stored_path: str) -> Path:
    path = Path(stored_path)
    return path if path.is_absolute() else root / path


def unique_terms(terms: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for term in terms:
        key = term.lower()
        if key not in seen:
            seen.add(key)
            result.append(term)
    return result


def href_between(output_dir: Path, target: Path) -> str:
    relative = os.path.relpath(target.resolve(), start=output_dir.resolve())
    return quote(Path(relative).as_posix(), safe="/.-_~")


def html_id(value: str) -> str:
    return quote(value, safe="")


def escape(value: str) -> str:
    return html.escape(value, quote=False)


def escape_attr(value: str) -> str:
    return html.escape(value, quote=True)


def selected_materials(
    root: Path,
    courses: list[Course],
    records: list[MaterialRecord],
    course_name: str | None,
) -> list[tuple[Path, Course | None]]:
    selected_course = find_course(courses, course_name) if course_name else None
    if course_name and selected_course is None:
        return []
    targets: list[tuple[Path, Course | None]] = []
    target_courses = [selected_course] if selected_course else courses
    for course in target_courses:
        if course is None:
            continue
        for material in course_materials(root, course, records):
            targets.append((material, course))
    return targets


def create_summary(
    root: Path,
    material: Path,
    course: Course | None,
    text: str,
    model: str,
    local_only: bool,
    api_provider: str = "openai",
) -> SummaryRecord:
    if not local_only:
        instructions = (
            "あなたは大学生の授業資料を分かりやすく整理・解説する優秀な学習支援AIアシスタントです。\n"
            "入力された授業資料を日本語で論理的に分析し、JSON形式でのみ回答を返してください。マークダウンのコードブロックなどで囲まず、純粋なJSON文字列だけを出力してください。\n"
            "JSONの構造は以下のキーを必ず持たせてください：\n"
            "1. `title`: 授業資料の内容を的確に表す、簡潔で分かりやすいタイトル（20文字程度）。\n"
            "2. `summary`: この授業回の「核心テーマ（何のための授業か、何を解決するのか）」から書き始め、授業全体の論理の流れが専門用語を知らない人にも直感的に伝わるように、300〜500文字程度で分かりやすく説明した文章。\n"
            "3. `key_points`: 授業の重要なポイントをまとめた文字列の配列。各要素は単一の文章として独立して意味が通り、具体的な内容を含むもの（例：「〜の定義とメリット」「〜における3つの主要課題」など）にしてください。\n"
            "4. `important_terms`: 学習上、絶対に外せない最重要キーワードとその簡単な解説を組み合わせた文字列の配列（例：「キーワード: ○○について解説した説明文」形式）。\n"
            "5. `review_checklist`: 授業内容が定着したかを自己確認（セルフテスト）するためのチェックリスト。単に「○○について」ではなく、「○○の仕組みを説明できるか？」「○○と○○の違いを3つ挙げられるか？」といった、具体的な疑問文形式の文字列の配列にしてください。\n"
        )
        user_text = (
            f"授業名: {course.name if course else 'その他'}\n"
            f"教員名: {course.teacher if course else ''}\n"
            f"資料名: {material.name}\n\n"
            f"資料本文:\n{trim_text(text)}"
        )
        if api_provider == "gemini" and os.environ.get("GEMINI_API_KEY"):
            try:
                payload = call_gemini_json(
                    model=model,
                    instructions=instructions,
                    user_text=user_text,
                )
                return SummaryRecord(
                    path=normalize_record_path(relative_to_root(root, material)),
                    course=course.name if course else None,
                    title=str(payload.get("title") or material.stem),
                    summary=str(payload.get("summary") or ""),
                    key_points=string_list(payload.get("key_points")),
                    important_terms=string_list(payload.get("important_terms")),
                    review_checklist=string_list(payload.get("review_checklist")),
                    source=f"gemini:{model}",
                )
            except Exception as exc:
                print(f"Gemini分析に失敗したためローカル要約に切り替えます: {exc}", file=sys.stderr)
        elif api_provider == "openai" and os.environ.get("OPENAI_API_KEY"):
            try:
                payload = call_openai_json(
                    model=model,
                    instructions=instructions,
                    user_text=user_text,
                )
                return SummaryRecord(
                    path=normalize_record_path(relative_to_root(root, material)),
                    course=course.name if course else None,
                    title=str(payload.get("title") or material.stem),
                    summary=str(payload.get("summary") or ""),
                    key_points=string_list(payload.get("key_points")),
                    important_terms=string_list(payload.get("important_terms")),
                    review_checklist=string_list(payload.get("review_checklist")),
                    source=f"openai:{model}",
                )
            except Exception as exc:
                print(f"OpenAI分析に失敗したためローカル要約に切り替えます: {exc}", file=sys.stderr)

    return local_summary(root, material, course, text)


def create_questions(
    course: Course,
    context: str,
    source_paths: list[str],
    count: int,
    model: str,
    local_only: bool,
    api_provider: str = "openai",
) -> list[QuestionRecord]:
    if not local_only:
        instructions = (
            "あなたは授業内容の定着度を測る問題作成AIです。"
            "入力された授業資料だけを根拠に、日本語で練習問題を作ってください。"
            "JSONだけを返してください。キーは questions です。"
            "questions は配列で、各要素は title, question_type, question, answer, explanation を持ちます。"
            "question_type は multiple_choice, short_answer, fill_blank のいずれかにしてください。"
        )
        user_text = (
            f"授業名: {course.name}\n"
            f"教員名: {course.teacher}\n"
            f"問題数: {count}\n\n"
            f"資料・要約:\n{trim_text(context)}"
        )
        if api_provider == "gemini" and os.environ.get("GEMINI_API_KEY"):
            try:
                payload = call_gemini_json(
                    model=model,
                    instructions=instructions,
                    user_text=user_text,
                )
                raw_questions = payload.get("questions", [])
                results: list[QuestionRecord] = []
                for item in raw_questions[:count]:
                    if not isinstance(item, dict):
                        continue
                    results.append(
                        QuestionRecord(
                            course=course.name,
                            title=str(item.get("title") or "練習問題"),
                            question_type=str(item.get("question_type") or "short_answer"),
                            question=str(item.get("question") or ""),
                            answer=str(item.get("answer") or ""),
                            explanation=str(item.get("explanation") or ""),
                            source_paths=source_paths,
                            source=f"gemini:{model}",
                        )
                    )
                if results:
                    return results
            except Exception as exc:
                print(f"Gemini問題生成に失敗したためローカル生成に切り替えます: {exc}", file=sys.stderr)
        elif api_provider == "openai" and os.environ.get("OPENAI_API_KEY"):
            try:
                payload = call_openai_json(
                    model=model,
                    instructions=instructions,
                    user_text=user_text,
                )
                raw_questions = payload.get("questions", [])
                results: list[QuestionRecord] = []
                for item in raw_questions[:count]:
                    if not isinstance(item, dict):
                        continue
                    results.append(
                        QuestionRecord(
                            course=course.name,
                            title=str(item.get("title") or "練習問題"),
                            question_type=str(item.get("question_type") or "short_answer"),
                            question=str(item.get("question") or ""),
                            answer=str(item.get("answer") or ""),
                            explanation=str(item.get("explanation") or ""),
                            source_paths=source_paths,
                            source=f"openai:{model}",
                        )
                    )
                if results:
                    return results
            except Exception as exc:
                print(f"OpenAI問題生成に失敗したためローカル生成に切り替えます: {exc}", file=sys.stderr)

    return local_questions(course, context, source_paths, count)


def call_openai_json(model: str, instructions: str, user_text: str) -> dict:
    try:
        from openai import OpenAI
    except ImportError as exc:
        raise RuntimeError("openai パッケージがインストールされていません。") from exc

    client = OpenAI()
    response = client.responses.create(
        model=model,
        instructions=instructions,
        input=user_text,
        text={"format": {"type": "json_object"}},
    )
    output_text = getattr(response, "output_text", "")
    if not output_text:
        raise RuntimeError("AI応答からテキストを取得できませんでした。")
    return json.loads(output_text)


def call_gemini_json(model: str, instructions: str, user_text: str) -> dict:
    try:
        from google import genai
        from google.genai import types
    except ImportError:
        import subprocess
        import sys
        print("Installing google-genai...", file=sys.stderr)
        try:
            subprocess.run([sys.executable, "-m", "pip", "install", "google-genai"], check=True, stdout=subprocess.DEVNULL)
            from google import genai
            from google.genai import types
        except Exception as e:
            raise RuntimeError(f"google-genai パッケージのインストールに失敗しました: {e}")

    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("環境変数 GEMINI_API_KEY が設定されていません。")

    client = genai.Client(api_key=api_key)
    response = client.models.generate_content(
        model=model,
        contents=user_text,
        config=types.GenerateContentConfig(
            system_instruction=instructions,
            response_mime_type="application/json",
        ),
    )
    if not response.text:
        raise RuntimeError("Geminiからの応答が空でした。")
    
    text = response.text.strip()
    if text.startswith("```"):
        lines = text.splitlines()
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines[-1].startswith("```"):
            lines = lines[:-1]
        text = "\n".join(lines).strip()
        
    return json.loads(text)


def local_summary(
    root: Path, material: Path, course: Course | None, text: str
) -> SummaryRecord:
    sentences = split_sentences(text)
    terms = [term for term, _count in Counter(tokenize(text)).most_common(10)]
    key_points = sentences[:5] if sentences else terms[:5]
    summary = " ".join(sentences[:3]) if sentences else " ".join(terms[:8])
    checklist = [f"{term} を説明できる" for term in terms[:5]]
    return SummaryRecord(
        path=normalize_record_path(relative_to_root(root, material)),
        course=course.name if course else None,
        title=material.stem,
        summary=summary or "本文から十分な要約を作成できませんでした。",
        key_points=key_points,
        important_terms=terms,
        review_checklist=checklist,
        source="local",
    )


def local_questions(
    course: Course, context: str, source_paths: list[str], count: int
) -> list[QuestionRecord]:
    terms = [term for term, _count in Counter(tokenize(context)).most_common(max(count, 1) * 2)]
    sentences = split_sentences(context)
    questions: list[QuestionRecord] = []
    for index in range(count):
        term = terms[index % len(terms)] if terms else course.name
        sentence = sentences[index % len(sentences)] if sentences else context[:120]
        questions.append(
            QuestionRecord(
                course=course.name,
                title=f"{course.name} 確認問題 {index + 1}",
                question_type="short_answer",
                question=f"「{term}」について、授業資料の内容に基づいて説明してください。",
                answer=term,
                explanation=sentence or "資料本文を見直して、関連する説明を確認してください。",
                source_paths=source_paths,
                source="local",
            )
        )
    return questions


def practice_context(
    root: Path,
    course: Course,
    records: list[MaterialRecord],
    summaries: list[SummaryRecord],
) -> tuple[str, list[str]]:
    course_summaries = [summary for summary in summaries if summary.course == course.name]
    if course_summaries:
        parts = []
        source_paths = []
        for summary in course_summaries:
            source_paths.append(summary.path)
            parts.append(
                "\n".join(
                    [
                        f"資料: {summary.title}",
                        f"要約: {summary.summary}",
                        "要点: " + " / ".join(summary.key_points),
                        "重要語句: " + " / ".join(summary.important_terms),
                    ]
                )
            )
        return "\n\n".join(parts), source_paths

    materials = course_materials(root, course, records)
    parts = []
    source_paths = []
    for material in materials:
        text = extract_text(material, "jpn+eng")
        if text.strip():
            source_paths.append(normalize_record_path(relative_to_root(root, material)))
            parts.append(f"資料: {material.name}\n{text}")
    return trim_text("\n\n".join(parts)), source_paths


def find_summary(
    summaries: list[SummaryRecord], root: Path, material: Path
) -> SummaryRecord | None:
    relative = normalize_record_path(relative_to_root(root, material))
    for summary in summaries:
        if normalize_record_path(summary.path) == relative:
            return summary
    return None


def upsert_summary(summaries: list[SummaryRecord], summary: SummaryRecord) -> None:
    summaries[:] = [
        item
        for item in summaries
        if normalize_record_path(item.path) != normalize_record_path(summary.path)
    ]
    summaries.append(summary)


def write_summary_pdf(root: Path, material: Path, summary: SummaryRecord) -> None:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    output_pdf = material.with_suffix(material.suffix + ".summary.pdf")
    
    font_path = "C:\\Windows\\Fonts\\msgothic.ttc"
    if not os.path.exists(font_path):
        font_path = "C:\\Windows\\Fonts\\meiryo.ttc"
        
    font_name = "MSGothic"
    try:
        pdfmetrics.registerFont(TTFont(font_name, font_path))
    except Exception as e:
        print(f"Warning: Failed to load Japanese font: {e}", file=sys.stderr)
        font_name = "Helvetica"
        
    doc = SimpleDocTemplate(
        str(output_pdf.resolve()),
        pagesize=A4,
        leftMargin=40,
        rightMargin=40,
        topMargin=40,
        bottomMargin=40
    )
    
    content_width = A4[0] - 80 # 515.27 pt
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#0f766e"),
        spaceAfter=12,
        wordWrap='CJK'
    )
    
    meta_style = ParagraphStyle(
        'MetaText',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#4b5563"),
        wordWrap='CJK'
    )
    
    heading_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=12,
        leading=16,
        textColor=colors.HexColor("#0f766e"),
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True,
        wordWrap='CJK'
    )
    
    body_style = ParagraphStyle(
        'BodyText',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        leading=15,
        textColor=colors.HexColor("#1f2937"),
        spaceAfter=8,
        wordWrap='CJK'
    )
    
    list_item_style = ParagraphStyle(
        'ListItemText',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        leading=15,
        textColor=colors.HexColor("#1f2937"),
        wordWrap='CJK'
    )

    story = []
    
    import html
    story.append(Paragraph(html.escape(summary.title), title_style))
    
    meta_data = [
        [Paragraph(f"<b>授業:</b> {html.escape(summary.course or 'その他')}", meta_style)],
        [Paragraph(f"<b>元資料:</b> {html.escape(str(summary.path))}", meta_style)],
        [Paragraph(f"<b>生成元:</b> {html.escape(summary.source)}", meta_style)]
    ]
    meta_table = Table(meta_data, colWidths=[content_width])
    meta_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#f9fafb")),
        ('PADDING', (0, 0), (-1, -1), 6),
        ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 15))
    
    story.append(Paragraph("【要約】", heading_style))
    story.append(Paragraph(html.escape(summary.summary), body_style))
    story.append(Spacer(1, 8))
    
    if summary.key_points:
        story.append(Paragraph("【要点】", heading_style))
        points_data = []
        for point in summary.key_points:
            points_data.append([
                Paragraph("<font color='#0f766e'>•</font>", list_item_style),
                Paragraph(html.escape(point), list_item_style)
            ])
        points_table = Table(points_data, colWidths=[12, content_width - 12])
        points_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
        ]))
        story.append(points_table)
        story.append(Spacer(1, 8))
        
    if summary.important_terms:
        story.append(Paragraph("【重要語句】", heading_style))
        terms_data = []
        for term in summary.important_terms:
            terms_data.append([
                Paragraph("<font color='#0f766e'>•</font>", list_item_style),
                Paragraph(html.escape(term), list_item_style)
            ])
        terms_table = Table(terms_data, colWidths=[12, content_width - 12])
        terms_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
        ]))
        story.append(terms_table)
        story.append(Spacer(1, 8))
        
    if summary.review_checklist:
        story.append(Paragraph("【復習チェックリスト】", heading_style))
        checklist_data = []
        for item in summary.review_checklist:
            checklist_data.append([
                Paragraph("<font color='#0f766e'>□</font>", list_item_style),
                Paragraph(html.escape(item), list_item_style)
            ])
        checklist_table = Table(checklist_data, colWidths=[15, content_width - 15])
        checklist_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
        ]))
        story.append(checklist_table)
        
    doc.build(story)


def write_practice_markdown(
    root: Path, course: Course, questions: list[QuestionRecord]
) -> None:
    output = root / course.folder / "practice_questions.md"
    lines = [f"# {course.name} 練習問題", ""]
    for index, question in enumerate(questions, start=1):
        lines.extend(
            [
                f"## {index}. {question.title}",
                "",
                f"- 種別: {question.question_type}",
                f"- 生成元: {question.source}",
                "",
                question.question,
                "",
                "### 解答",
                "",
                question.answer,
                "",
                "### 解説",
                "",
                question.explanation,
                "",
            ]
        )
    output.write_text("\n".join(lines), encoding="utf-8")


def split_sentences(text: str) -> list[str]:
    compact = re.sub(r"\s+", " ", text).strip()
    parts = re.split(r"(?<=[。！？.!?])\s*", compact)
    return [part.strip() for part in parts if len(part.strip()) >= 12]


def trim_text(text: str, limit: int = MAX_AI_CHARS) -> str:
    text = text.strip()
    if len(text) <= limit:
        return text
    return text[:limit] + "\n\n[...省略...]"


def string_list(value: object) -> list[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, str):
        return [item.strip() for item in re.split(r"[,、\n]", value) if item.strip()]
    return []


def recalculate_course_learned_terms(course: Course, records: list[MaterialRecord]) -> None:
    course_records = [r for r in records if r.course == course.name]
    if not course_records:
        course.learned_terms = []
        return

    term_doc_counts = Counter()
    for r in course_records:
        for term in r.learned_terms:
            if tokenize(term):
                term_doc_counts[term] += 1

    course_keywords_tokens = set(tokenize(" ".join(course.keywords)))
    
    sorted_terms = []
    for term, count in term_doc_counts.most_common():
        if term not in course_keywords_tokens:
            sorted_terms.append(term)
            
    course.learned_terms = sorted_terms[:80]


def find_course(courses: list[Course], name: str | None) -> Course | None:
    if not name:
        return None
    for course in courses:
        if course.name == name or course.folder == name:
            return course
    return None


def relative_to_root(root: Path, path: Path) -> str:
    try:
        return str(path.resolve().relative_to(root))
    except ValueError:
        return str(path.resolve())


def normalize_record_path(path: str) -> str:
    return path.replace("\\", "/")


def find_material_record(
    records: list[MaterialRecord], root: Path, material: Path
) -> MaterialRecord | None:
    relative = normalize_record_path(relative_to_root(root, material))
    absolute = normalize_record_path(str(material.resolve()))
    for record in records:
        stored = normalize_record_path(record.path)
        if stored == relative or stored == absolute:
            return record
    return None


def infer_course_from_path(root: Path, courses: list[Course], material: Path) -> str | None:
    relative_parent = normalize_record_path(relative_to_root(root, material.parent))
    for course in courses:
        if normalize_record_path(course.folder) == relative_parent:
            return course.name
    return None


def upsert_material_record(
    records: list[MaterialRecord],
    root: Path,
    path: Path,
    course: str | None,
    learned_terms: list[str],
    score: float,
) -> None:
    relative = normalize_record_path(relative_to_root(root, path))
    records[:] = [
        record
        for record in records
        if normalize_record_path(record.path) != relative
    ]
    records.append(
        MaterialRecord(
            path=relative,
            course=course,
            learned_terms=learned_terms,
            score=score,
        )
    )


def unique_destination(path: Path) -> Path:
    if not path.exists():
        return path
    stem = path.stem
    suffix = path.suffix
    for index in range(1, 1000):
        candidate = path.with_name(f"{stem}_{index}{suffix}")
        if not candidate.exists():
            return candidate
    raise RuntimeError(f"保存先ファイル名を決められません: {path}")


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="授業資料を読み取り、授業ごとのフォルダへ自動分類します。"
    )
    parser.add_argument(
        "--root",
        help="授業フォルダと設定を保存するルート。省略時は現在のフォルダ。",
    )
    subparsers = parser.add_subparsers(dest="command", required=True)

    add = subparsers.add_parser("add-course", help="授業を登録して保存フォルダを作成")
    add.add_argument("name", help="授業名")
    add.add_argument("--teacher", default="", help="教員名")
    add.add_argument("--keywords", default="", help="カンマ区切りの授業キーワード")
    add.set_defaults(func=add_course)

    update = subparsers.add_parser("update-course", help="登録済み授業のキーワードや教員名を変更")
    update.add_argument("name", help="変更したい授業の授業名")
    update.add_argument("--teacher", help="新しい教員名（省略時は変更なし）")
    update.add_argument("--keywords", help="新しいキーワード（カンマ区切り、省略時は変更なし）")
    update.set_defaults(func=update_course)

    listing = subparsers.add_parser("list-courses", help="登録済み授業の一覧を表示")
    listing.set_defaults(func=list_courses)

    sorter = subparsers.add_parser("sort", help="資料を授業ごとのフォルダへ分類")
    sorter.add_argument("inputs", nargs="+", help="資料ファイルまたは資料フォルダ")
    sorter.add_argument("--recursive", action="store_true", help="フォルダを再帰的に処理")
    sorter.add_argument("--copy", action="store_true", help="移動せずコピーする")
    sorter.add_argument(
        "--threshold",
        type=float,
        default=0.08,
        help="分類に必要なスコア。低いほど分類されやすい。",
    )
    sorter.add_argument(
        "--ocr-language",
        default="jpn+eng",
        help="Tesseract OCR の言語指定。例: jpn, eng, jpn+eng",
    )
    sorter.set_defaults(func=sort_materials)

    reassign = subparsers.add_parser(
        "reassign", help="誤分類された資料を別の授業フォルダへ移動"
    )
    reassign.add_argument("material", help="移動したい資料ファイル")
    reassign.add_argument("--to", required=True, help="正しい移動先の授業名")
    reassign.add_argument(
        "--ocr-language",
        default="jpn+eng",
        help="履歴がない資料を移動する場合のOCR言語指定。",
    )
    reassign.set_defaults(func=reassign_material)

    delete_parser = subparsers.add_parser("delete-material", help="資料を削除し、関連する語句も削除")
    delete_parser.add_argument("material", help="削除したい資料ファイル")
    delete_parser.add_argument(
        "--ocr-language",
        default="jpn+eng",
        help="履歴がない資料を削除する場合のOCR言語指定。",
    )
    delete_parser.set_defaults(func=delete_material)

    analyze = subparsers.add_parser(
        "analyze", help="資料をAIまたはローカル処理で分析し、要約と要点を生成"
    )
    analyze.add_argument("--course", help="対象の授業名。省略時は全授業。")
    analyze.add_argument(
        "--model",
        default=DEFAULT_MODEL,
        help=f"使用するモデル。省略時は {DEFAULT_MODEL}。",
    )
    analyze.add_argument(
        "--api-provider",
        choices=["openai", "gemini"],
        default="openai",
        help="使用するAIプロバイダー（openai または gemini）。",
    )
    analyze.add_argument(
        "--local-only",
        action="store_true",
        help="APIを使わず、抽出型のローカル要約だけを行う。",
    )
    analyze.add_argument("--force", action="store_true", help="既存の要約を再生成する。")
    analyze.add_argument(
        "--ocr-language",
        default="jpn+eng",
        help="Tesseract OCR の言語指定。例: jpn, eng, jpn+eng",
    )
    analyze.set_defaults(func=analyze_materials)

    practice = subparsers.add_parser(
        "generate-practice", help="授業資料から練習問題を生成"
    )
    practice.add_argument("--course", help="対象の授業名。省略時は全授業。")
    practice.add_argument("--count", type=int, default=5, help="授業ごとの問題数。")
    practice.add_argument(
        "--model",
        default=DEFAULT_MODEL,
        help=f"使用するモデル。省略時は {DEFAULT_MODEL}。",
    )
    practice.add_argument(
        "--api-provider",
        choices=["openai", "gemini"],
        default="openai",
        help="使用するAIプロバイダー（openai または gemini）。",
    )
    practice.add_argument(
        "--local-only",
        action="store_true",
        help="APIを使わず、語句ベースのローカル問題を生成する。",
    )
    practice.add_argument("--force", action="store_true", help="既存問題を置き換える。")
    practice.set_defaults(func=generate_practice)

    site = subparsers.add_parser(
        "build-site", help="授業資料をブラウザで見やすいHTMLページにまとめる"
    )
    site.add_argument(
        "--output",
        default="_site",
        help="HTMLを書き出すフォルダ。省略時は _site。",
    )
    site.set_defaults(func=build_site)

    clear = subparsers.add_parser(
        "clear-analysis", help="作成した要約や練習問題を消去する"
    )
    clear.add_argument("--course", help="対象の授業名。省略時は全授業。")
    clear.add_argument("--material", help="対象の資料ファイル。省略時は全資料。")
    clear.set_defaults(func=clear_analysis)

    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    return args.func(args)


if __name__ == "__main__":
    raise SystemExit(main())
